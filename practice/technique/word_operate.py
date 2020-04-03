"""
 !/usr/bin/env python3.6
 -*- coding: utf-8 -*-
 --------------------------------
 Description : word相关的文件操作
 参考链接：
 https://www.cnblogs.com/ontheway703/p/5266041.html
 https://blog.csdn.net/xtfge0915/article/details/83479922
 --------------------------------
 @Time    : 2019/5/25 10:39
 @File    : word_operate.py
 @Software: PyCharm
 --------------------------------
 @Author  : lixj
 @contact : lixj_zj@163.com
"""

from docx import *
import re
from docx.shared import Pt
from docx.shared import Inches
from win32com import client
from docx.enum.style import WD_STYLE_TYPE
import os, zipfile, shutil


class WordOperate():
    def __init__(self):
        pass

    def show_style(self):
        """
        查看word常用样式，在读取或写入时设置
        :return:
        """
        doc = Document()
        styles = doc.styles

        # 查看段落样式
        for s in styles:
            if s.type == WD_STYLE_TYPE.PARAGRAPH:
                print("段落样式：", s.name)

        # 查看字符样式
        for s in styles:
            if s.type == WD_STYLE_TYPE.CHARACTER:
                print("字符样式：", s.name)

        # 查看表格样式
        for s in styles:
            if s.type == WD_STYLE_TYPE.TABLE:
                print("表格样式：", s.name)

    def set_style(self):
        """
        word中设置样式
        :return:
        """
        document = Document()

        # 1.段落设置样式
        paragraph = document.add_paragraph()
        paragraph.style = document.styles['Heading 1']  # style选取
        paragraph.style = 'Heading 1'  # 用样式名称直接赋值
        paragraph2 = document.add_paragraph(style='Body Text')  # 创建段落时赋值

        # 2.设置段落中的字符格式，定义样式中的字符格式后，所有运用此样式的段落都有相应的字符格式
        #   从样式库中选取 'Normal' 样式，并提取 'Normal' 样式的字符属性
        style = document.styles['Normal']
        font = style.font
        #   设置样式中的字符属性 ，操作方法和上面改变内联对象属性方法一致
        font.name = "Microsoft YaHei UI"
        font.size = Pt(50)  # 字体大小
        #   将设置好字符属性的样式运用到段落中
        p = document.add_paragraph("change font attribution", style='Normal')

        # 3.设置段落格式，定义样式中的段落格式后，所有运用此样式的段落都有相应的段落格式
        styles = document.styles
        #   选取 style，并设置 style 中的段落格式
        style = styles['Heading 2']
        para_format = style.paragraph_format
        para_format.left_indent = Pt(20)
        para_format.widow_control = True
        #   将设置好段落格式的 style 运用到段落中
        p = document.add_paragraph('This is Heading, level 1', style=style)

    def get_head(self, doc_path):
        """
        读取 word 标题
        :param doc_path: 文档路径
        :return:
        """
        doc = Document(doc_path)
        for p in doc.paragraphs:
            # 遍历1、2、3级标题
            if p.style.name == 'Heading 1':
                print(p.text)
            if p.style.name == 'Heading 2':
                print(p.text)
            if p.style.name == 'Heading 3':
                print(p.text)

            # 遍历所有标题
            for p in doc.paragraphs:
                if re.match("^Heading \d+$", p.style.name):
                    print(p.text)

    def get_content(self, doc_path):
        """
        读取 word 内容
        :param doc_path: 文档路径
        :return:
        """
        doc = Document(doc_path)
        # 读取正文
        for p in doc.paragraphs:
            if p.style.name == 'Normal':
                print(p.text)

    def write_content(self, result_doc_path):
        """
        doc写入文件
        :param result_doc_path: 写入 word 文档路径
        :return:
        """
        doc = Document()
        # 写入标题
        doc.add_heading("heading 1", level=1)
        doc.add_paragraph("heading 1", style='Heading 1')

        # 写入正文
        doc.add_paragraph("正文")

        # 写入分页符
        doc.add_page_break()

        # 写入表格
        table = doc.add_table(rows=1, cols=3, style="Light List Accent 5")
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'testName'
        hdr_cells[1].text = 'param'
        hdr_cells[2].text = 'exc'

        # 写入图片
        doc.add_picture("imgName", width=Inches(1.5))  # 设置宽度

        doc.save(result_doc_path)

    def get_img_from_doc(self, doc_path):
        """
        从word中提取图片到对应目录img文件夹
        word转zip-提取midea中的word-复制到新目录并重命名同名为文件夹-zip还原成word，删除word文件夹
        :param docdir: word 路径
        :return:
        """
        # 切换路径
        doc_abs_path = os.path.abspath(os.path.dirname(doc_path) + os.path.sep + ".")
        os.chdir(doc_abs_path)

        # 遍历文件
        for file in os.listdir(doc_abs_path):
            if file.endswith("docx"):  # 匹配docx文件
                doc_name = file.split(".")  # 以“.”做成列表形式
                os.rename(file, "{docName}.ZIP".format(docName=doc_name[0]))  # 重命名为ZIP格式
                f = zipfile.ZipFile("{docName}.ZIP".format(docName=doc_name[0]), 'r')
                for file in f.namelist():
                    if "word" in file:
                        f.extract(file)  # 将压缩包里的word文件夹解压
                f.close()
                old_img_dir = r"{absPath}\word\media".format(absPath=doc_abs_path)  # 定义图片文件夹
                shutil.copytree(old_img_dir, "{absPath}\{docName}".format(absPath=doc_abs_path,
                                                                          docName=doc_name[0]))  # 拷贝到新目录，名称为word文件的名字
                os.rename("{docName}.ZIP".format(docName=doc_name[0]),
                          "{docName}.docx".format(docName=doc_name[0]))  # 将ZIP名字还原为DOCX
                shutil.rmtree("{absPath}\word".format(absPath=doc_abs_path))  # 删除word文件夹
            else:
                print(file, "非docx文件！")

    def word_to_pdf(self, doc_path):
        """
        word 转换成 pdf
        :param doc_path: docx 文件路径
        :return:
        """
        # 切换路径
        doc_abs_path = os.path.abspath(os.path.dirname(doc_path) + os.path.sep + ".")
        os.chdir(doc_abs_path)
        dir_list = os.listdir(doc_abs_path)
        for file in dir_list:
            if file.endswith("docx"):  # 匹配docx文件
                word = client.DispatchEx("Word.Application")
                base_name = os.path.basename(file).split('.')[0]
                word_file = word.Documents.Open(os.path.abspath(file), ReadOnly=1)
                word_file.SaveAs(doc_abs_path + os.path.sep + str(base_name) + ".pdf", FileFormat=17)
                word_file.Close()
            else:
                print(file, "非docx文件！")


if __name__ == '__main__':
    root_doc_path = r'D:\ZX\temp\test\1\test.docx'
    result_doc_path = r"C:\Users\Tebon\Desktop\test\result.docx"
    word_oper = WordOperate()
    word_oper.get_content(root_doc_path)
    word_oper.write_content(result_doc_path)
    word_oper.show_style()
    word_oper.get_img_from_doc(root_doc_path)
    word_oper.word_to_pdf(root_doc_path)
